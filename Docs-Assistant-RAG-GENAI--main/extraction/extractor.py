# extraction/extractor.py
"""
Extraction module for medical document assistant.
Handles PDF, DOCX, and PNG.
Returns structured dict: { file_name, file_type, pages: [...] }
"""

import io
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import pytesseract
import cv2
import numpy as np


# ---------- Schema Helpers ----------

def page_record(page_number: int, text: str = "", tables=None):
    return {
        "page_number": page_number,
        "text": text or "",
        "tables": [{"rows": tbl} for tbl in (tables or [])],
        "images": []
    }


def extraction_response(file_name: str, file_type: str, pages):
    return {
        "file_name": file_name,
        "file_type": file_type,
        "pages": pages
    }


# ---------- PDF Extraction ----------

def extract_pdf(file_bytes: bytes, file_name: str = "uploaded.pdf"):
    pages_out = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            raw_tables = page.extract_tables() or []
            clean_tables = []
            for tbl in raw_tables:
                rows = []
                for row in tbl:
                    rows.append([("" if cell is None else str(cell)).strip() for cell in row])
                clean_tables.append(rows)
            pages_out.append(page_record(i, text, clean_tables))
    return extraction_response(file_name, "pdf", pages_out)


# ---------- DOCX Extraction ----------

def extract_docx(file_bytes: bytes, file_name: str = "uploaded.docx"):
    doc = DocxDocument(io.BytesIO(file_bytes))
    text_blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    doc_text = "\n".join(text_blocks)

    clean_tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            cells = [(c.text or "").strip() for c in r.cells]
            rows.append(cells)
        clean_tables.append(rows)

    pages_out = [page_record(1, doc_text, clean_tables)]
    return extraction_response(file_name, "docx", pages_out)


# ---------- PNG Extraction (OCR + Table Detection) ----------

def ocr_image_to_text(pil_img: Image.Image) -> str:
    return pytesseract.image_to_string(pil_img) or ""


def detect_table_from_image(pil_img: Image.Image):
    img = np.array(pil_img.convert("RGB"))
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    thr = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                cv2.THRESH_BINARY_INV, 15, 10)

    horizontalkernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
    detect_horizontal = cv2.morphologyEx(thr, cv2.MORPH_OPEN, horizontalkernel, iterations=2)

    verticalkernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 25))
    detect_vertical = cv2.morphologyEx(thr, cv2.MORPH_OPEN, verticalkernel, iterations=2)

    grid = cv2.addWeighted(detect_horizontal, 0.5, detect_vertical, 0.5, 0.0)
    grid = cv2.dilate(grid, np.ones((3, 3), np.uint8), iterations=1)

    contours, _ = cv2.findContours(grid, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
    boxes = []
    H, W = gray.shape
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if w > W // 100 and h > H // 100:
            boxes.append((x, y, w, h))

    if not boxes:
        return []

    boxes = sorted(boxes, key=lambda b: (b[1], b[0]))
    rows = []
    current_row = [boxes[0]]
    for b in boxes[1:]:
        if abs(b[1] - current_row[-1][1]) < 15:
            current_row.append(b)
        else:
            rows.append(sorted(current_row, key=lambda r: r[0]))
            current_row = [b]
    rows.append(sorted(current_row, key=lambda r: r[0]))

    table = []
    for r in rows:
        row_texts = []
        for (x, y, w, h) in r:
            crop = pil_img.crop((x, y, x + w, y + h))
            txt = pytesseract.image_to_string(crop, config="--psm 6")
            row_texts.append(txt.strip())
        table.append(row_texts)

    return table


def extract_png(file_bytes: bytes, file_name: str = "uploaded.png"):
    pil_img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    img = np.array(pil_img)

    # Grayscale & threshold
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    binary = cv2.adaptiveThreshold(~gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C,
                                   cv2.THRESH_BINARY, 15, -2)

    # Detect horizontal
    horizontal = binary.copy()
    cols = horizontal.shape[1]
    horizontal_size = max(10, cols // 30)   # smaller kernel
    horizontalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (horizontal_size, 1))
    horizontal = cv2.erode(horizontal, horizontalStructure)
    horizontal = cv2.dilate(horizontal, horizontalStructure)

    # Detect vertical
    vertical = binary.copy()
    rows = vertical.shape[0]
    vertical_size = max(10, rows // 30)
    verticalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (1, vertical_size))
    vertical = cv2.erode(vertical, verticalStructure)
    vertical = cv2.dilate(vertical, verticalStructure)

    # Combine
    mask = cv2.add(horizontal, vertical)

    # Find contours
    contours, _ = cv2.findContours(mask, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
    boxes = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if w > 10 and h > 10:   # relaxed condition
            boxes.append((x, y, w, h))

    table_data = []
    if boxes:
        # Sort by rows
        boxes = sorted(boxes, key=lambda b: (b[1], b[0]))
        rows = []
        current_row = []
        last_y = None
        for b in boxes:
            x, y, w, h = b
            if last_y is None or abs(y - last_y) < 20:  # allow more variation
                current_row.append(b)
                last_y = y
            else:
                rows.append(sorted(current_row, key=lambda r: r[0]))
                current_row = [b]
                last_y = y
        if current_row:
            rows.append(sorted(current_row, key=lambda r: r[0]))

        # OCR per cell
        for r in rows:
            row_texts = []
            for (x, y, w, h) in r:
                crop = pil_img.crop((x, y, x + w, y + h))
                txt = pytesseract.image_to_string(crop, config="--psm 6")
                row_texts.append(txt.strip())
            table_data.append(row_texts)

    # Always fallback to full text if no structured table
    full_text = pytesseract.image_to_string(pil_img)

    pages_out = [page_record(1, full_text if not table_data else "", 
                             [table_data] if table_data else [])]
    return extraction_response(file_name, "png", pages_out)


# ---------- Dispatcher ----------

def extract_file(file_bytes: bytes, file_name: str):
    if file_name.lower().endswith(".pdf"):
        return extract_pdf(file_bytes, file_name)
    elif file_name.lower().endswith(".docx"):
        return extract_docx(file_bytes, file_name)
    elif file_name.lower().endswith(".png"):
        return extract_png(file_bytes, file_name)
    else:
        raise ValueError("Unsupported file type. Only PDF, DOCX, PNG allowed.")
